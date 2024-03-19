
from openai_utils import get_completion_openai, init_openai
from prompt_utils import build_prompt
from pdf_utils import  extract_text_from_pdf_pdfplumber_with_pages
from text_utils import split_text, split_text_with_pages
from vectordb_utils_shule import ShuleVectorDB
from sentence_transformers import CrossEncoder
from dotenv import load_dotenv, find_dotenv
_ = load_dotenv(find_dotenv()) 
import os
import pandas as pd
import time
import sys
import errno
import pickle
from logger import log_info, log_debug, log_warning
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt

top_n = 30
show_n = 8
recall_n = 120
select_questions = [0, 1, 2] # None for all questions
countries = ["South Asia"]

search_strategy = "rerank"
vec_db_shule = None
rerank_model = CrossEncoder(os.getenv('RERANK_MODEL_PATH')) if search_strategy == "rerank" else None

prompt_template_doc = """
You are a scientific Q&A bot with expertise in antimicrobial resistance, one health, environmental science and policy making. You answer user question based on the information provided by the user above the question and your in-house knowledge. There are five pieces of extra information above the user question. You answer in uses question's language. The user question is in the final line. When you use the user information, always indicate the Reference in your answer. Additionally, let us know which part of your answer is from the user's information and which part is based on your in-house knowledge by writing either [Reference] or [In-house knowledge]. 

__INFO__

## User's question：
__QUERY__

"""

def init_db_load_index(index_path):
    global vec_db_shule
    init_openai()
    log_info("---init database by index file begin---")
    log_info(f"index path: {index_path}")
    with open(index_path, 'rb') as file:
        vec_db_shule = pickle.load(file)
    
    log_info("---init database by index file end---")
    
def rerank(user_input, top_n, recall_n):
    search_labels = vec_db_shule.search_bge(user_input, recall_n)
    t0 = time.time()
    texts, pages, titles, years, countries, ORGs = vec_db_shule.get_context_by_labels(search_labels)
    t1 = time.time()
    log_info(f"vec_db_shule.get_context_by_labels costs: {t1 - t0}")

    documents = [texts[i] if countries[i] == 'xxx' else f'In {countries[i]}, {texts[i]}' for i in range(len(pages))]
    res = rerank_model.rank(documents = documents,
                            query=user_input,
                            batch_size = 1,
                            return_documents = False)
    t2 = time.time()
    log_info(f"rerank_model.predict costs: {t2 - t1}")
    
    ids = [i['corpus_id'] for i in res][:top_n]
    scores = [i['score'] for i in res][:top_n]

    # sorted_list = {'scores': scores, 
    #                 'texts': [texts[i] for i in ids], 
    #                 'titles':[titles[i] for i in ids], 
    #                 'years':[years[i] for i in ids], 
    #                 'countries':[countries[i] for i in ids], 
    #                 'pages':[pages[i] for i in ids], 
    #                 'ORGs':[ORGs[i] for i in ids]}
    log_info(f"finish rerank {recall_n} texts, return highest {top_n} texts")
    # for score, doc in sorted_list:
    #     print(f"{score}\t{doc}\n")
    return scores, [texts[i] for i in ids], [pages[i] for i in ids], [titles[i] for i in ids], [years[i] for i in ids], [countries[i] for i in ids], [ORGs[i] for i in ids]


def build_doc_for_country(country: str):
    log_info(f"{country} begins...")
    t0=time.time()
    document = Document()
    document.styles["Normal"].font.name = "Times New Roman"
    document.styles["Normal"].font.size = Pt(12) 
    # document.styles["Title"].font.name = "Times New Roman"
    # document.styles['Heading 1'].font.name = "Times New Roman"
    document.styles['Intense Quote'].font.name = "Times New Roman"
    document.styles['Intense Quote'].font.size = Pt(14)
    
    indicator_path = os.getenv('INDICATOR_PATH')
    AMR_indicators = pd.read_csv(indicator_path)
    first_page = True
    
    for i in range(AMR_indicators.shape[0]):
        if select_questions != None and i not in select_questions: continue
        if first_page:
            first_page = False
        else:
            document.add_page_break()
            
        indictor = AMR_indicators.iloc[i]
        question = f'{indictor["Question"][:-1]} in {country}?'
        # document.add_heading(f'{indictor["Panel"]}', 0)
        document.add_heading(level = 0).add_run(f'{indictor["Panel"]}').font.name = "Times New Roman"
        document.add_heading(level = 0).add_run(f'{indictor["Domain"]}').font.name = "Times New Roman"
        p = document.add_heading(level = 1)
        p.add_run(f'{indictor["Number"]}: {question}').font.name = "Times New Roman"
        p.paragraph_format.space_after = Pt(6) 
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            
        scores, texts, pages, titles, years, countries, ORGs = rerank(question, top_n, recall_n)
        
        res = [texts[i] if countries[i] == 'xxx' else f'In {countries[i]}, {texts[i]}' for i in range(top_n)]
        # search_field = "\n\n".join([f"{i+1}. [Reference: {titles[i]}, Page: {pages[i]}, ORG: {ORGs[i]}, Year: {years[i]}]\n{texts[i]}" for i in range(top_n)])
        prompt = build_prompt(template=prompt_template_doc, info=[f"{res[i]} [Reference: Page {pages[i]}, {titles[i]}, {years[i]}, {ORGs[i]}]" for i in range(top_n)], query=question)
        response = get_completion_openai(prompt, [])
        log_info(response)
        for line in response.split("\n"):
            p = document.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6) 
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        document.add_paragraph()
        
        p_ref = document.add_paragraph('Reference', style='Intense Quote')
        p_ref.paragraph_format.left_indent = 0
        for i in range(show_n):
            p = document.add_paragraph(f'{ORGs[i]}. ({years[i]}). ')
            p.add_run(titles[i]).italic = True
            p.add_run(f". Page {pages[i]}.")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(6) 

            p2 = document.add_paragraph(texts[i])
            p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p2.paragraph_format.space_after = Pt(6) 

            document.add_paragraph().paragraph_format.space_after = 0
            
        document.save(f'./generated_docs/Evaluation_HandBook_{country}.docx')
    log_info(f"{country} costs: {time.time() - t0}")
                
def main():
    for country in countries:
        build_doc_for_country(country)
    
    
if __name__ == "__main__":
    init_db_load_index(index_path=os.getenv('INDEX_PATH'))
    main()